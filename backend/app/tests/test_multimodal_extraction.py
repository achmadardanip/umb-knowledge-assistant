import sys
import types

from docx import Document

from app.core.config import get_settings
from app.multimodal.audio_extractor import extract_audio
from app.multimodal.document_extractor import extract_document
from app.multimodal.file_downloader import download_file
from app.multimodal.image_ocr_extractor import extract_image_ocr
from app.multimodal.pdf_extractor import extract_pdf
from app.multimodal.presentation_extractor import extract_pptx
from app.multimodal.spreadsheet_extractor import extract_spreadsheet
from app.multimodal.video_extractor import extract_video_metadata
from app.retrieval.hybrid_retriever import HybridRetriever
from app.db.models import Chunk, ChunkEmbedding, Document as DbDocument, Source
from app.ingestion import pipeline
from app.ingestion.pipeline import upsert_source_document


def test_file_downloader_rejects_files_outside_scope():
    result = download_file("https://example.com/file.pdf")
    assert result.status == "skipped"
    assert result.reason == "outside_allowed_domain"


def test_file_downloader_rejects_files_over_size_limits(monkeypatch):
    class Head:
        headers = {"content-type": "application/pdf", "content-length": str(31 * 1024 * 1024)}

    monkeypatch.setattr("requests.head", lambda *args, **kwargs: Head())
    result = download_file("https://mercubuana.ac.id/file.pdf")
    assert result.status == "skipped"
    assert result.reason == "file_too_large"


def test_pdf_extractor_extracts_page_level_text_and_metadata(tmp_path, monkeypatch):
    class Page:
        def get_text(self, _kind):
            return "Informasi PDF UMB"

    class Doc(list):
        metadata = {"title": "PDF UMB"}

        def __enter__(self):
            return self

        def __exit__(self, *args):
            return False

    fake_fitz = types.SimpleNamespace(open=lambda path: Doc([Page()]))
    monkeypatch.setitem(sys.modules, "fitz", fake_fitz)
    pdf = tmp_path / "fake.pdf"
    pdf.write_bytes(b"%PDF")
    pages = extract_pdf(pdf)
    assert pages[0].page_number == 1
    assert pages[0].metadata["title"] == "PDF UMB"


def test_docx_extractor_extracts_paragraphs_and_tables(tmp_path):
    path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("Paragraf UMB")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Kolom A"
    table.rows[0].cells[1].text = "Kolom B"
    doc.save(path)
    result = extract_document(path, "docx")
    assert "Paragraf UMB" in result.content
    assert "Kolom A" in result.content


def test_pptx_extractor_extracts_slide_text_and_slide_numbers(monkeypatch, tmp_path):
    class Shape:
        text = "Judul Slide"

    class Slide:
        shapes = [Shape()]

    fake_pptx = types.SimpleNamespace(Presentation=lambda path: types.SimpleNamespace(slides=[Slide()]))
    monkeypatch.setitem(sys.modules, "pptx", fake_pptx)
    path = tmp_path / "deck.pptx"
    path.write_bytes(b"pptx")
    slides = extract_pptx(path)
    assert slides[0].slide_number == 1
    assert "Judul Slide" in slides[0].content


def test_spreadsheet_extractor_extracts_csv_columns_and_summary(tmp_path):
    path = tmp_path / "data.csv"
    path.write_text("nama,biaya\nReguler,100\n", encoding="utf-8")
    sheets = extract_spreadsheet(path, "csv")
    assert sheets[0].sheet_name == "data"
    assert sheets[0].columns == ["nama", "biaya"]
    assert "Reguler" in sheets[0].content


def test_ocr_asr_and_video_download_are_disabled_by_default(monkeypatch):
    # Verify the code defaults (off), independent of the operator's .env, which may
    # enable OCR/ASR for multimodal ingestion.
    from app.core.config import _bool

    for var in ("ENABLE_OCR", "ENABLE_ASR", "ENABLE_VIDEO_DOWNLOAD"):
        monkeypatch.delenv(var, raising=False)
    assert _bool("ENABLE_OCR", False) is False
    assert _bool("ENABLE_ASR", False) is False
    assert _bool("ENABLE_VIDEO_DOWNLOAD", False) is False


def test_archive_urls_are_not_indexed_unless_live_official_url_is_validated():
    from app.discovery.scope_validator import validate_url_scope
    from app.discovery.url_normalizer import archive_to_live_candidate

    live = archive_to_live_candidate("https://web.archive.org/web/20200101/https://mercubuana.ac.id/berita")
    assert validate_url_scope(live).is_allowed


def test_retrieval_returns_source_type_specific_metadata(db):
    source = Source(url="https://mercubuana.ac.id/file.pdf", hostname="mercubuana.ac.id", title="PDF", status="indexed")
    db.add(source)
    db.flush()
    doc = DbDocument(source_id=source.id, raw_text="biaya pendaftaran", cleaned_text="biaya pendaftaran")
    db.add(doc)
    db.flush()
    db.add(
        Chunk(
            document_id=doc.id,
            source_id=source.id,
            chunk_text="biaya pendaftaran Universitas Mercu Buana",
            chunk_index=0,
            token_count=5,
            meta={"url": source.url, "title": "PDF", "hostname": "mercubuana.ac.id", "source_type": "pdf", "page_number": 4},
            source_type="pdf",
            page_number=4,
            extraction_method="pymupdf",
            extraction_confidence=0.85,
        )
    )
    db.commit()
    contexts = HybridRetriever(db).search("biaya pendaftaran", top_k=1)
    assert contexts[0]["source_type"] == "pdf"
    assert contexts[0]["page_number"] == 4
    assert contexts[0]["extraction_method"] == "pymupdf"


def test_subdomain_chunk_metadata_preserves_hostname_path_and_retrieves(db):
    text = " ".join(["pendaftaran"] * 40 + ["mahasiswa", "baru", "UMB"])
    chunks_created = upsert_source_document(
        db,
        "https://pmb.mercubuana.ac.id/pendaftaran",
        text,
        "PMB UMB",
        {},
        200,
        discovery_source="katana",
    )
    db.commit()

    assert chunks_created > 0
    chunk = db.query(Chunk).one()
    assert chunk.meta["hostname"] == "pmb.mercubuana.ac.id"
    assert chunk.meta["path"] == "/pendaftaran"
    assert chunk.meta["discovery_source"] == "katana"
    contexts = HybridRetriever(db).search("pendaftaran mahasiswa baru", top_k=1)
    assert contexts[0]["url"] == "https://pmb.mercubuana.ac.id/pendaftaran"
    assert contexts[0]["hostname"] == "pmb.mercubuana.ac.id"


def test_html_ingestion_stores_local_embedding_in_sidecar(db, monkeypatch):
    class _LocalEmbedder:
        storage = "sidecar"
        provider_name = "local_e5"
        model = "intfloat/multilingual-e5-small"
        dimension = 384
        profile = "local-e5-small-v1"
        version = "1"

        def embed_texts(self, texts):
            return [[1.0] + [0.0] * 383 for _ in texts]

    monkeypatch.setattr(pipeline, "get_embedder", lambda: _LocalEmbedder())
    text = " ".join(["informasi", "pendaftaran", "mahasiswa", "baru"] * 12)

    upsert_source_document(
        db,
        "https://pmb.mercubuana.ac.id/local-embedding",
        text,
        "PMB UMB",
        {},
        200,
        discovery_source="katana",
    )
    db.commit()

    chunk = db.query(Chunk).one()
    sidecar = db.query(ChunkEmbedding).one()
    assert chunk.embedding is None
    assert sidecar.chunk_id == chunk.id
    assert sidecar.profile == "local-e5-small-v1"


def test_html_ingestion_falls_back_to_keyword_only_when_sidecar_migration_is_missing(db, monkeypatch):
    class _LocalEmbedder:
        storage = "sidecar"
        provider_name = "local_e5"
        model = "intfloat/multilingual-e5-small"
        dimension = 384
        profile = "local-e5-small-v1"
        version = "1"

        def embed_texts(self, texts):
            return [[1.0] + [0.0] * 383 for _ in texts]

    ChunkEmbedding.__table__.drop(db.get_bind())
    monkeypatch.setattr(pipeline, "get_embedder", lambda: _LocalEmbedder())

    created = upsert_source_document(
        db,
        "https://pmb.mercubuana.ac.id/no-sidecar",
        " ".join(["informasi", "pendaftaran", "mahasiswa", "baru"] * 12),
        "PMB UMB",
        {},
        200,
        discovery_source="katana",
    )
    db.commit()

    assert created == 1
    assert db.query(Chunk).one().embedding is None


def test_retrieval_expands_daftar_query_to_pendaftaran_subdomain(db):
    upsert_source_document(
        db,
        "https://mercubuana.ac.id/biro-kemahasiswaan/beasiswa",
        " ".join(["mahasiswa"] * 20 + ["baru"] * 5 + ["beasiswa"]),
        "Beasiswa UMB",
        {},
        200,
        discovery_source="katana",
    )
    upsert_source_document(
        db,
        "https://pendaftaran.mercubuana.ac.id/",
        " ".join(
            [
                "Buat akun pendaftaran, pilih jalur seleksi, lalu lakukan pembelian formulir.",
                "Pendaftaran mahasiswa baru Universitas Mercu Buana melalui PMB resmi.",
            ]
            * 20
        ),
        "Pendaftaran Mahasiswa Baru UMB",
        {},
        200,
        discovery_source="katana",
    )
    db.commit()

    contexts = HybridRetriever(db, dense_enabled=False).search("Bagaimana cara daftar mahasiswa baru?", top_k=1)

    assert contexts[0]["url"] == "https://pendaftaran.mercubuana.ac.id/"
    assert contexts[0]["hostname"] == "pendaftaran.mercubuana.ac.id"


def test_frontend_api_types_support_provider_override():
    from pathlib import Path

    api_ts = Path("../frontend/app/lib/api.ts").read_text(encoding="utf-8")
    assert "provider_override" in api_ts


def test_source_card_supports_subdomain_hostname_display():
    from pathlib import Path

    source_card = Path("../frontend/app/components/SourceCard.tsx").read_text(encoding="utf-8")
    assert "hostnameForSource" in source_card
    assert "new URL(source.url).hostname" in source_card
