from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


@dataclass
class ExtractedDocument:
    content: str
    tables: list[str]
    extraction_method: str
    extraction_confidence: float = 0.85


def extract_docx(path: str | Path) -> ExtractedDocument:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables: list[str] = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            tables.append(f"Tabel {table_index}\n" + "\n".join(rows))
    return ExtractedDocument(
        content="\n".join(parts + tables).strip(),
        tables=tables,
        extraction_method="python-docx",
    )


def extract_document(path: str | Path, source_type: str) -> ExtractedDocument:
    if source_type == "docx":
        return extract_docx(path)
    return ExtractedDocument(
        content="",
        tables=[],
        extraction_method="unsupported-doc-conversion",
        extraction_confidence=0.0,
    )

